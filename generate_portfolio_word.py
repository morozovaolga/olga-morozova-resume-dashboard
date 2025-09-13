#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для создания Word документа портфолио Ольги Морозовой
на основе данных из лендинга
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
import os
from datetime import datetime

class PortfolioWordGenerator:
    def __init__(self, filename="Olga_Morozova_Portfolio.docx"):
        self.filename = filename
        self.doc = Document()
        
        # Настройка базовых стилей
        self.setup_styles()
        
        # Данные портфолио (те же, что в PDF версии)
        self.portfolio_data = {
            'name': 'Ольга Морозова',
            'title': 'Главный редактор, контент-стратег и руководитель проектов',
            'description': 'Помогаю компаниям создавать сильный контент, который работает на рост аудитории и бизнеса.',
            
            'contacts': {
                'phone': '+7 (916) 841-32-16',
                'email': 'morozova31@gmail.com',
                'telegram': '@oljona',
                'location': 'Москва / Удалённо'
            },
            
            'about': {
                'description': 'Я руководитель цифровых знаний с более чем 10-летним опытом создания и управления цифровыми библиотеками, контент-стратегиями и образовательными проектами.',
                'experience': 'Специализируюсь на организации информационных систем, создании удобных интерфейсов и разработке контент-стратегий, которые делают знания доступными для всех.',
                'points': [
                    'Системный подход: Строю процессы, автоматизирую рутину, навожу порядок в информационном хаосе.',
                    'Ориентация на результат: Работаю на измеримый результат и пользу для бизнеса.',
                    'Командная работа: Умею строить и мотивировать команды для сложных проектов.',
                    'Успешные проекты: Запустила и вела десятки цифровых продуктов и сервисов.'
                ]
            },
            
            'experience': {
                'company': 'Российская государственная библиотека',
                'position': 'Главный редактор мобильного приложения «Свет»',
                'period': 'Февраль 2022 — Январь 2025',
                'location': 'Москва',
                'description': 'Возглавляла продуктовую команду, отвечающую за контент, аналитику и развитие цифрового сервиса для миллионов пользователей.',
                'achievements': [
                    'Увеличила KPI (книговыдачи) в 3 раза (до 2,5-3 млн в год) за счет стратегии контента, партнерств и аналитики',
                    'Разработала и координировала 17 крупных партнерских проектов',
                    'Сотрудничество с Московской электронной школой — 3 млн потенциальных пользователей, 100 тыс. книговыдач в 2024',
                    'Акция «Лето с Пятерочкой» — охват 1350 магазинов, рост переходов на сайт в 10 раз за месяц',
                    'Построила систему сбора и визуализации данных (Яндекс.Метрика, Superset, Excel)',
                    'Создала внутренние сервисы для автоматизации — сокращение времени подготовки текстов с 2 дней до 30 минут',
                    'Перевела команду на российский инструмент (Kaiten), создав систему контроля процессов',
                    'Разработала 11 лендингов в Tilda',
                    'Управляла редакцией из 5 человек'
                ]
            },
            
            'skills': {
                'professional': [
                    'Управление проектами',
                    'Контент-стратегия', 
                    'Продуктовое мышление',
                    'Аналитика данных',
                    'Системное мышление',
                    'Управление командой'
                ],
                'technical': [
                    'Цифровые библиотеки',
                    'Пользовательский опыт',
                    'Информационная архитектура',
                    'Цифровая трансформация',
                    'Автоматизация',
                    'Аналитика'
                ],
                'languages': [
                    'Русский (родной)',
                    'Английский (свободно)',
                    'Французский (разговорный)'
                ]
            },
            
            'education': {
                'degrees': [
                    {
                        'degree': 'Магистр журналистики',
                        'institution': 'Московский государственный университет им. М.В. Ломоносова',
                        'program': 'Международная программа: Mastère franco-russe de journalisme'
                    },
                    {
                        'degree': 'Филология',
                        'institution': 'Новосибирский национальный исследовательский государственный университет'
                    }
                ],
                'certifications': [
                    'Аналитик данных',
                    'Продуктовая аналитика',
                    'Гибкая разработка (Agile) с Jira',
                    'Авторское право (ВОИС)',
                    'Инструменты анализа данных (Яндекс Практикум)',
                    'Менеджер проектов (Яндекс Практикум)',
                    'Контент-маркетинг (Нетология)'
                ]
            },
            
            'cases': [
                {
                    'title': 'Аудиокниги на борту Аэрофлота',
                    'background': 'Аэрофлот хотел улучшить впечатления пассажиров в полете с помощью цифрового контента. В рамках партнерства с РГБ мы разработали мультимедийный сервис аудиокниг.',
                    'role': [
                        'Курировала контент: отобрала и организовала коллекцию аудиокниг',
                        'Разработала стандарты дизайна: создала обложки с использованием ИИ',
                        'Координировала интеграцию контента для приложения и лендинга',
                        'Взаимодействовала с техническими командами'
                    ],
                    'results': [
                        'Успешно запустили новую функцию для пассажиров Аэрофлота',
                        'Проект представлен Министру культуры и получил одобрение',
                        'Укрепили позиционирование как новатора',
                        'Продемонстрировали потенциал ИИ-дизайна в культурных проектах'
                    ]
                },
                {
                    'title': 'Интеграция с Московской электронной школой',
                    'background': 'МЭШ — крупнейшая образовательная платформа России. РГБ интегрировала свой каталог для обеспечения доступа к литературе школьной программы.',
                    'role': [
                        'Координировала интеграцию библиотеки с МЭШ через API',
                        'Отбирала каталог книг по школьным программам',
                        'Контролировала подготовку контента и метаданных',
                        'Была связующим звеном между командами'
                    ],
                    'results': [
                        'Интегрированы тысячи книг в экосистему МЭШ',
                        'Обеспечен доступ для 1+ млн учеников и учителей',
                        'Повышена образовательная ценность платформы',
                        'Укреплена роль библиотеки как лидера в образовании'
                    ]
                },
                {
                    'title': 'Московское метро «Подземный читальный зал»',
                    'background': 'Московский метрополитен хотел вовлечь пассажиров в культурные инициативы. Совместно с РГБ превратили станцию в интерактивный читальный зал через QR-коды.',
                    'role': [
                        'Разработала концепцию контента и курировала библиотеку',
                        'Спроектировала механику доступа через QR-коды',
                        'Координировала с властями метро и техническими командами',
                        'Контролировала коммуникационные материалы'
                    ],
                    'results': [
                        'Запущена первая интерактивная культурная инициатива в метро',
                        'Превращен транспортный узел в брендированный опыт',
                        'Привлечено внимание СМИ и положительные отзывы',
                        'Укреплен имидж библиотеки как новатора'
                    ]
                },
                {
                    'title': 'Проект к юбилею Александра Пушкина',
                    'background': 'К юбилею Пушкина РГБ инициировала масштабный гибридный проект для продвижения классической литературы в современном формате.',
                    'role': [
                        'Разработала концепцию контента и структуру викторины',
                        'Создала контент в разных форматах: викторины, стикеры, раскраски',
                        'Координировала участие 140 библиотек в 50 регионах',
                        'Руководила онлайн и офлайн распространением'
                    ],
                    'results': [
                        'Успешное общенациональное мероприятие в 50 регионах',
                        'Тысячи участников через онлайн и офлайн каналы',
                        'Повышена видимость классической литературы',
                        'Признание профессионального сообщества'
                    ]
                }
            ],
            
            'portfolio': {
                'landing_pages': [
                    'Книги о счастливой любви',
                    'Читаем в небе',
                    '100 лет Расулу Гамзатову',
                    'Новогодние предсказания',
                    'Квиз к 225-летию А.С. Пушкина'
                ],
                'design_materials': [
                    'Раскраски к юбилею А.С. Пушкина',
                    'Новогодние украшения с символикой проекта',
                    'Инфографика для библиотек',
                    'Макеты брендированных иллюстраций',
                    'Стикеры для Telegram'
                ],
                'documentation': [
                    'Руководство для мобильных библиотек 2024',
                    'Инструкции для партнерских библиотек'
                ],
                'texts': [
                    {
                        'title': 'Дни культуры Таджикистана в РГБ',
                        'subtitle': 'Репортаж с открытия двух выставок, приуроченных к Дням культуры Республики Таджикистан в России.',
                        'url': 'https://www.rsl.ru/ru/events/afisha/vistavki/dni-kultury-tadzhikistana-v-rgb'
                    },
                    {
                        'title': 'XVIII Международные научные чтения памяти Николая Фёдоровича Фёдорова',
                        'subtitle': 'Торжественное открытие XVIII Международных научных чтений памяти философа Н. Ф. Фёдорова.',
                        'url': 'https://www.rsl.ru/ru/all-news/xviii-fedorovskie-chtenia'
                    },
                    {
                        'title': 'Круглый стол «Уникальный лингвистический фонд коренных народов cеверных территорий — мировая сокровищница культуры»',
                        'subtitle': 'Круглый стол, посвящённый Международному году языков коренных народов Севера.',
                        'url': 'https://www.rsl.ru/ru/events/afisha/meetings/kruglyij-stol-unikalnyij-lingvisticheskij-fond'
                    },
                    {
                        'title': 'Круглый стол «День ООН в РГБ — 2019»',
                        'subtitle': 'Круглый стол, посвящённый Дню ООН и 74-й сессии Генеральной Ассамблеи ООН.',
                        'url': 'https://www.rsl.ru/ru/events/afisha/meetings/kruglyij-stol-den-oon-v-rgb-2019'
                    },
                    {
                        'title': 'В РГБ открылась фотовыставка пейзажей Южного берега Крыма',
                        'subtitle': 'Торжественное открытие баннерной выставки «Берег мечты. Культурное и природное наследие Южного берега Крыма».',
                        'url': 'https://www.rsl.ru/ru/all-news/fotovyistavka-pejzazhej-yuzhnogo-berega-kryima'
                    },
                    {
                        'title': 'Понаснимали тут',
                        'subtitle': 'Стартовал "анонимный" этап конкурса на лучший фоторепортаж о Москве "Серебряная камера-2005".',
                        'url': 'https://web.archive.org/web/20061130092311/http://lenta.ru:80/articles/2006/10/03/foto/'
                    },
                    {
                        'title': 'И уэльбекерно, и тошно',
                        'subtitle': 'Скандальный французский писатель привез в Россию новый роман',
                        'url': 'https://web.archive.org/web/20120214221335/http://lenta.ru:80/articles/2006/04/20/houellebecq/'
                    },
                    {
                        'title': '"Надеюсь, вы не купили мою книгу"',
                        'subtitle': 'Фредерик Бегбедер прилетел на свидание с Москвой',
                        'url': 'https://web.archive.org/web/20120708064328/http://lenta.ru/articles/2006/07/03/beigbeder/'
                    },
                    {
                        'title': 'Свидание с Ктулху',
                        'subtitle': 'Роман про морских чудищ захватил книжные витрины',
                        'url': 'https://disk.yandex.ru/i/5dekldkfLa54dA'
                    }
                ]
            }
        }
    
    def setup_styles(self):
        """Настройка стилей документа"""
        styles = self.doc.styles
        
        # Создаем кастомные стили
        try:
            # Стиль заголовка
            title_style = styles.add_style('CustomTitle', 1)  # WD_STYLE_TYPE.PARAGRAPH
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(51, 76, 153)  # Синий цвет
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Стиль заголовка секции
            section_style = styles.add_style('SectionTitle', 1)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(16)
            section_font.bold = True
            section_font.color.rgb = RGBColor(51, 76, 153)
            
            # Стиль подзаголовка
            subsection_style = styles.add_style('SubsectionTitle', 1)
            subsection_font = subsection_style.font
            subsection_font.name = 'Calibri'
            subsection_font.size = Pt(14)
            subsection_font.bold = True
            subsection_font.color.rgb = RGBColor(76, 76, 76)
            
        except ValueError:
            # Стили уже существуют
            pass
    
    def add_header(self):
        """Добавление заголовка документа"""
        title = self.doc.add_paragraph(self.portfolio_data['name'], 'CustomTitle')
        subtitle = self.doc.add_paragraph(self.portfolio_data['title'], 'SubsectionTitle')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        description = self.doc.add_paragraph(self.portfolio_data['description'])
        description.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем разделитель
        self.doc.add_paragraph()
    
    def add_contacts(self):
        """Добавление контактной информации"""
        self.doc.add_paragraph('Контактная информация', 'SectionTitle')
        
        contacts = self.portfolio_data['contacts']
        contact_para = self.doc.add_paragraph()
        contact_para.add_run(f"Телефон: ").bold = True
        contact_para.add_run(f"{contacts['phone']}\n")
        contact_para.add_run(f"Email: ").bold = True
        contact_para.add_run(f"{contacts['email']}\n")
        contact_para.add_run(f"Telegram: ").bold = True
        contact_para.add_run(f"{contacts['telegram']}\n")
        contact_para.add_run(f"Местоположение: ").bold = True
        contact_para.add_run(f"{contacts['location']}")
        
        self.doc.add_paragraph()
    
    def add_about_section(self):
        """Добавление секции О себе"""
        self.doc.add_paragraph('О себе', 'SectionTitle')
        
        about = self.portfolio_data['about']
        self.doc.add_paragraph(about['description'])
        self.doc.add_paragraph(about['experience'])
        
        self.doc.add_paragraph('Ключевые качества:', 'SubsectionTitle')
        for point in about['points']:
            p = self.doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(point)
        
        self.doc.add_paragraph()
    
    def add_experience_section(self):
        """Добавление секции Опыт работы"""
        self.doc.add_paragraph('Профессиональный опыт', 'SectionTitle')
        
        exp = self.portfolio_data['experience']
        
        # Заголовок позиции
        position_para = self.doc.add_paragraph()
        position_para.add_run(exp['position']).bold = True
        position_para.add_run(f"\n{exp['company']}")
        position_para.add_run(f"\n{exp['period']} • {exp['location']}")
        
        # Описание
        self.doc.add_paragraph(exp['description'])
        
        # Достижения
        self.doc.add_paragraph('Ключевые достижения:', 'SubsectionTitle')
        for achievement in exp['achievements']:
            p = self.doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(achievement)
        
        self.doc.add_paragraph()
    
    def add_skills_section(self):
        """Добавление секции Навыки"""
        self.doc.add_paragraph('Ключевые компетенции', 'SectionTitle')
        
        skills = self.portfolio_data['skills']
        
        # Профессиональные навыки
        self.doc.add_paragraph('Профессиональные навыки:', 'SubsectionTitle')
        skills_para = self.doc.add_paragraph(' • '.join(skills['professional']))
        
        # Технические компетенции
        self.doc.add_paragraph('Технические компетенции:', 'SubsectionTitle')
        self.doc.add_paragraph(' • '.join(skills['technical']))
        
        # Языки
        self.doc.add_paragraph('Языковые навыки:', 'SubsectionTitle')
        self.doc.add_paragraph(' • '.join(skills['languages']))
        
        self.doc.add_paragraph()
    
    def add_education_section(self):
        """Добавление секции Образование"""
        self.doc.add_paragraph('Образование и повышение квалификации', 'SectionTitle')
        
        education = self.portfolio_data['education']
        
        # Основное образование
        self.doc.add_paragraph('Образование:', 'SubsectionTitle')
        for degree in education['degrees']:
            degree_para = self.doc.add_paragraph()
            degree_para.add_run(degree['degree']).bold = True
            degree_para.add_run(f"\n{degree['institution']}")
            if 'program' in degree:
                degree_para.add_run(f"\nПрограмма: {degree['program']}")
            self.doc.add_paragraph()
        
        # Сертификация
        self.doc.add_paragraph('Повышение квалификации:', 'SubsectionTitle')
        for cert in education['certifications']:
            p = self.doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(cert)
        
        self.doc.add_paragraph()
    
    def add_cases_section(self):
        """Добавление секции Кейсы"""
        # Добавляем разрыв страницы
        self.doc.add_page_break()
        
        self.doc.add_paragraph('Кейсы', 'SectionTitle')
        self.doc.add_paragraph('Реальные проекты с измеримыми результатами и бизнес-эффектом')
        self.doc.add_paragraph()
        
        for i, case in enumerate(self.portfolio_data['cases']):
            # Заголовок кейса
            self.doc.add_paragraph(f"{i+1}. {case['title']}", 'SubsectionTitle')
            
            # Задача
            task_para = self.doc.add_paragraph()
            task_para.add_run('Задача: ').bold = True
            task_para.add_run(case['background'])
            
            # Моя роль
            role_para = self.doc.add_paragraph()
            role_para.add_run('Моя роль:').bold = True
            for role_item in case['role']:
                p = self.doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(role_item)
            
            # Результаты
            results_para = self.doc.add_paragraph()
            results_para.add_run('Результаты:').bold = True
            for result in case['results']:
                p = self.doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(result)
            
            self.doc.add_paragraph()
    
    def add_portfolio_section(self):
        """Добавление секции Портфолио"""
        self.doc.add_paragraph('Портфолио', 'SectionTitle')
        self.doc.add_paragraph('Примеры проектов: от лендингов до дизайн-системы')
        
        portfolio = self.portfolio_data['portfolio']
        
        # Лендинги
        self.doc.add_paragraph('Лендинги спецпроектов:', 'SubsectionTitle')
        for project in portfolio['landing_pages']:
            p = self.doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(project)
        
        # Дизайн-материалы
        self.doc.add_paragraph('Дизайн-материалы:', 'SubsectionTitle')
        for material in portfolio['design_materials']:
            p = self.doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(material)
        
        # Документация
        self.doc.add_paragraph('Документация и руководства:', 'SubsectionTitle')
        for doc in portfolio['documentation']:
            p = self.doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(doc)
        
        # Тексты
        self.doc.add_paragraph('Тексты:', 'SubsectionTitle')
        for text in portfolio['texts']:
            p = self.doc.add_paragraph()
            p.add_run('• ').bold = True
            title_run = p.add_run(text['title'])
            title_run.bold = True
            
            p.add_run(f"\n  {text['subtitle']}")
            
            # Теперь все тексты имеют ссылки
            url_run = p.add_run(f"\n  Ссылка: {text['url']}")
            url_run.font.color.rgb = RGBColor(51, 76, 153)  # Синий цвет для ссылки
            url_run.italic = True
        
        self.doc.add_paragraph()
    
    def add_footer(self):
        """Добавление подвала документа"""
        self.doc.add_paragraph()
        footer_para = self.doc.add_paragraph(f"Документ сгенерирован {datetime.now().strftime('%d.%m.%Y')}")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate_docx(self):
        """Генерация Word документа"""
        print("Создание содержимого Word документа...")
        self.add_header()
        self.add_contacts()
        self.add_about_section()
        self.add_experience_section()
        self.add_skills_section()
        self.add_education_section()
        self.add_cases_section()
        self.add_portfolio_section()
        self.add_footer()
        
        print(f"Сохранение Word файла: {self.filename}")
        self.doc.save(self.filename)
        print(f"✅ Word документ успешно создан: {self.filename}")
        
        return self.filename

def main():
    """Основная функция"""
    print("🔄 Запуск генерации портфолио в Word...")
    
    # Создание генератора и генерация Word документа
    generator = PortfolioWordGenerator("Olga_Morozova_Portfolio.docx")
    docx_filename = generator.generate_docx()
    
    # Проверка размера файла
    if os.path.exists(docx_filename):
        file_size = os.path.getsize(docx_filename)
        print(f"📄 Размер файла: {file_size / 1024:.1f} KB")
        print(f"📍 Путь к файлу: {os.path.abspath(docx_filename)}")
    
    print("✨ Готово! Word портфолио создано успешно.")

if __name__ == "__main__":
    main()
